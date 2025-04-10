# editor_report_writer.py
# Compiles structured editorial feedback into a beautifully formatted, actionable Word document report

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

# Use absolute path for output folder
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "output_reports")

# --- Write editorial report to DOCX ---
def write_report(filename, feedback_list, total_score, verdict):
    doc = Document()

    # Set title style
    doc.add_heading("📝 Editorial Feedback Report", level=0)

    # Metadata block
    intro = doc.add_paragraph()
    intro.add_run(f"File Reviewed: ").bold = True
    intro.add_run(f"{filename}\n")
    intro.add_run("Date: ").bold = True
    intro.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    intro.add_run("Verdict: ").bold = True
    intro.add_run(f"{verdict}\n")
    intro.add_run("Total Score: ").bold = True
    intro.add_run(f"{total_score} / 10\n")

    doc.add_paragraph("\n")

    # Each chunk's feedback, with page ranges and separation
    for i, (chunk_info, feedback) in enumerate(feedback_list):
        section_title = f"Chunk {i+1} (Pages {chunk_info['start_page']}-{chunk_info['end_page']})"
        heading = doc.add_heading(section_title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        para = doc.add_paragraph()
        para.add_run(feedback).font.size = Pt(11)

        # Horizontal line
        doc.add_paragraph("\u2014" * 50)

    # Ensure output folder exists
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    # Sanitize filename and build full output path
    base_name = os.path.splitext(os.path.basename(filename))[0]
    base_name_clean = "_".join(base_name.split())
    report_name = f"{base_name_clean}_feedback.docx"
    report_path = os.path.join(OUTPUT_FOLDER, report_name)

    print(f"📛 Base filename: {base_name_clean}")
    print(f"📝 Saving report to: {report_path}")

    try:
        doc.save(report_path)
        if os.path.exists(report_path):
            print("📂 Report saved successfully.")
        else:
            print("⚠️ File save attempted but no file found at target path.")
    except Exception as e:
        print(f"❌ Failed to save report: {e}")
        return None

    return report_path
